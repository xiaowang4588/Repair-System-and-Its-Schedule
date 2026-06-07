"""
报告渲染器 v2.0

专业级报告生成，领导视角：
  1. Excel 报告 - 完整数据表格，专业格式
  2. Word 报告 - 正式文档，可直接打印汇报
"""
import logging
from io import BytesIO
from datetime import datetime

logger = logging.getLogger(__name__)


class ReportRenderer:
    """报告渲染器"""

    def _create_font(self, bold=False, size=10, color='000000'):
        from openpyxl.styles import Font
        return Font(bold=bold, size=size, color=color)

    def _create_fill(self, color='4472C4'):
        from openpyxl.styles import PatternFill
        return PatternFill(start_color=color, end_color=color, fill_type='solid')

    def _create_border(self):
        from openpyxl.styles import Border, Side
        return Border(
            left=Side(style='thin', color='B0B0B0'),
            right=Side(style='thin', color='B0B0B0'),
            top=Side(style='thin', color='B0B0B0'),
            bottom=Side(style='thin', color='B0B0B0')
        )

    def _create_align(self, h='left', v='center'):
        from openpyxl.styles import Alignment
        return Alignment(horizontal=h, vertical=v, wrap_text=True)

    def _set_header(self, cell, value):
        """设置表头单元格"""
        cell.value = value
        cell.font = self._create_font(bold=True, size=10, color='FFFFFF')
        cell.fill = self._create_fill('2F5496')
        cell.alignment = self._create_align('center')
        cell.border = self._create_border()

    def _set_cell(self, cell, value, bold=False, align='left'):
        """设置普通单元格"""
        cell.value = value
        cell.font = self._create_font(bold=bold, size=10)
        cell.alignment = self._create_align(align)
        cell.border = self._create_border()

    def _set_title(self, cell, value):
        """设置标题"""
        cell.value = value
        cell.font = self._create_font(bold=True, size=14, color='2F5496')
        cell.alignment = self._create_align('center')

    def _set_subtitle(self, cell, value):
        """设置副标题"""
        cell.value = value
        cell.font = self._create_font(bold=True, size=11, color='2F5496')
        cell.alignment = self._create_align('left')

    # ============================================================
    # Excel 报告
    # ============================================================

    def render_excel(self, analysis: dict, advice: list = None) -> BytesIO:
        """渲染完整 Excel 报告"""
        from openpyxl import Workbook
        from openpyxl.utils import get_column_letter

        wb = Workbook()

        # Sheet 1: 报告概览
        self._excel_overview(wb, analysis)

        # Sheet 2: 报修明细（核心数据）
        self._excel_records_detail(wb, analysis)

        # Sheet 3: 按教室统计
        self._excel_by_classroom(wb, analysis)

        # Sheet 4: 按楼栋统计
        self._excel_by_building(wb, analysis)

        # Sheet 5: 按故障类型统计
        self._excel_by_fault_type(wb, analysis)

        # Sheet 6: 按处理人统计
        self._excel_by_handler(wb, analysis)

        # Sheet 7: 按学院统计
        self._excel_by_college(wb, analysis)

        # Sheet 8: 周对比分析（仅周报）
        if analysis.get('report_type') == 'weekly':
            self._excel_week_comparison(wb, analysis)

        # Sheet 9: 智能建议
        if advice:
            self._excel_advice(wb, advice)

        # === 新增 Sheet ===
        # Sheet 10: 多周趋势（仅周报）
        if analysis.get('report_type') == 'weekly':
            self._excel_multi_week_trend(wb, analysis)

        # Sheet 11: 设备健康度
        self._excel_equipment_health(wb, analysis)

        # Sheet 12: 楼栋画像
        self._excel_building_profile(wb, analysis)

        # Sheet 13: 响应时效
        self._excel_response_time(wb, analysis)

        # Sheet 14: 工作负载
        self._excel_workload(wb, analysis)

        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer

    def _excel_overview(self, wb, analysis: dict):
        """Sheet 1: 报告概览"""
        from openpyxl.utils import get_column_letter

        ws = wb.active
        ws.title = '报告概览'

        report_type = analysis.get('report_type', '')
        date_range = analysis.get('date_range', {})
        overview = analysis.get('overview', {})

        # 标题
        ws.merge_cells('A1:H1')
        self._set_title(ws['A1'], f'重庆移通学院綦江校区 多媒体设备运维{self._get_type_name(report_type)}')

        ws.merge_cells('A2:H2')
        ws['A2'].value = date_range.get('label', '')
        ws['A2'].font = self._create_font(size=11, color='666666')
        ws['A2'].alignment = self._create_align('center')

        ws.merge_cells('A3:H3')
        ws['A3'].value = f'报告生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M")}'
        ws['A3'].font = self._create_font(size=9, color='999999')
        ws['A3'].alignment = self._create_align('center')

        # 核心指标表格
        row = 5
        self._set_subtitle(ws.cell(row=row, column=1), '一、核心指标')
        row += 1

        # 表头
        headers = ['指标', '本周', '上周', '变化', '趋势']
        for col, h in enumerate(headers, 1):
            self._set_header(ws.cell(row=row, column=col), h)
        row += 1

        vs = overview.get('vs_last_week', {})
        metrics = [
            ('报修总量', overview.get('total_count', 0), vs.get('last_total', 0), vs.get('total_change', 0), vs.get('trend', '-')),
            ('已处理', overview.get('resolved_count', 0), vs.get('last_resolved', 0), vs.get('resolved_change', 0), ''),
            ('待处理', overview.get('pending_count', 0), vs.get('last_pending', 0), vs.get('pending_change', 0), ''),
            ('处理中', overview.get('processing_count', 0), '', '', ''),
            ('未处理', overview.get('unhandled_count', 0), '', '', ''),
            ('处理率', f"{overview.get('resolved_rate', 0)}%", f"{vs.get('last_rate', 0)}%", f"{vs.get('rate_change', 0):+.1f}%", ''),
            ('外聘教师报修', overview.get('external_teacher_count', 0), '', '', ''),
            ('设备更换次数', overview.get('device_replace_count', 0), '', '', ''),
            ('涉及教室数', overview.get('classroom_count', 0), '', '', ''),
            ('涉及楼栋数', overview.get('building_count', 0), '', '', ''),
            ('涉及学院数', overview.get('college_count', 0), '', '', ''),
        ]

        for metric_name, this_val, last_val, change, trend in metrics:
            self._set_cell(ws.cell(row=row, column=1), metric_name, bold=True)
            self._set_cell(ws.cell(row=row, column=2), this_val, align='center')
            self._set_cell(ws.cell(row=row, column=3), last_val, align='center')
            self._set_cell(ws.cell(row=row, column=4), change, align='center')
            self._set_cell(ws.cell(row=row, column=5), trend, align='center')
            row += 1

        # 设置列宽
        col_widths = [20, 15, 15, 15, 10]
        for i, w in enumerate(col_widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = w

    def _excel_records_detail(self, wb, analysis: dict):
        """Sheet 2: 报修明细"""
        from openpyxl.utils import get_column_letter

        ws = wb.create_sheet('报修明细')

        records = analysis.get('records_detail', [])

        # 标题
        ws.merge_cells('A1:N1')
        self._set_title(ws['A1'], '报修记录明细表')

        row = 3

        # 表头
        headers = [
            '序号', '报修日期', '星期', '故障教室', '报修类型', '故障原因',
            '处理状态', '报修人', '报修人学院', '是否外聘', '报修方式',
            '处理人', '是否换设备', '处理方案'
        ]
        for col, h in enumerate(headers, 1):
            self._set_header(ws.cell(row=row, column=col), h)
        row += 1

        # 数据行
        for i, r in enumerate(records, 1):
            values = [
                i,
                r.get('report_time', '')[:10],
                r.get('weekday', ''),
                r.get('classroom', ''),
                r.get('fault_type', ''),
                r.get('fault_cause', ''),
                r.get('status', ''),
                r.get('reporter_name', ''),
                r.get('reporter_college', ''),
                '是' if r.get('is_external_teacher') else '否',
                r.get('report_method', ''),
                r.get('handler_name', ''),
                '是' if r.get('is_device_replaced') else '否',
                r.get('solution', ''),
            ]
            for col, val in enumerate(values, 1):
                self._set_cell(ws.cell(row=row, column=col), val, align='center' if col <= 3 else 'left')
            row += 1

        # 设置列宽
        col_widths = [6, 12, 6, 12, 10, 30, 8, 10, 18, 8, 12, 10, 8, 30]
        for i, w in enumerate(col_widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = w

        # 冻结首行
        ws.freeze_panes = 'A4'

    def _excel_by_classroom(self, wb, analysis: dict):
        """Sheet 3: 按教室统计"""
        from openpyxl.utils import get_column_letter

        ws = wb.create_sheet('按教室统计')

        by_classroom = analysis.get('repairs_by_classroom', {})

        ws.merge_cells('A1:H1')
        self._set_title(ws['A1'], '按教室统计')

        row = 3

        headers = ['教室', '本周次数', '学期累计', '是否高频', '故障类型分布', '主要报修人', '主要处理人', '当前状态']
        for col, h in enumerate(headers, 1):
            self._set_header(ws.cell(row=row, column=col), h)
        row += 1

        for classroom, info in by_classroom.items():
            repairs = info.get('repairs', [])

            # 故障类型分布
            fault_types = {}
            reporters = {}
            handlers = {}
            statuses = {}

            for r in repairs:
                ft = r.get('fault_type', '')
                if ft:
                    fault_types[ft] = fault_types.get(ft, 0) + 1

                reporter = r.get('reporter', '')
                if reporter:
                    reporters[reporter] = reporters.get(reporter, 0) + 1

                handler = r.get('handler', '')
                if handler:
                    handlers[handler] = handlers.get(handler, 0) + 1

                status = r.get('status', '')
                if status:
                    statuses[status] = statuses.get(status, 0) + 1

            fault_str = '、'.join([f'{k}({v})' for k, v in sorted(fault_types.items(), key=lambda x: x[1], reverse=True)])
            top_reporter = max(reporters.items(), key=lambda x: x[1])[0] if reporters else ''
            top_handler = max(handlers.items(), key=lambda x: x[1])[0] if handlers else ''
            status_str = '、'.join([f'{k}({v})' for k, v in statuses.items()])

            self._set_cell(ws.cell(row=row, column=1), classroom, bold=True)
            self._set_cell(ws.cell(row=row, column=2), info.get('count', 0), align='center')
            self._set_cell(ws.cell(row=row, column=3), info.get('semester_total', 0), align='center')
            self._set_cell(ws.cell(row=row, column=4), '是' if info.get('is_frequent') else '否', align='center')
            self._set_cell(ws.cell(row=row, column=5), fault_str)
            self._set_cell(ws.cell(row=row, column=6), top_reporter)
            self._set_cell(ws.cell(row=row, column=7), top_handler)
            self._set_cell(ws.cell(row=row, column=8), status_str)
            row += 1

        col_widths = [15, 10, 10, 10, 30, 12, 12, 20]
        for i, w in enumerate(col_widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = w

    def _excel_by_building(self, wb, analysis: dict):
        """Sheet 4: 按楼栋统计"""
        from openpyxl.utils import get_column_letter

        ws = wb.create_sheet('按楼栋统计')

        by_building = analysis.get('repairs_by_building', {})

        ws.merge_cells('A1:F1')
        self._set_title(ws['A1'], '按楼栋统计')

        row = 3

        headers = ['楼栋', '报修次数', '占比', '故障类型分布', '已处理', '待处理']
        for col, h in enumerate(headers, 1):
            self._set_header(ws.cell(row=row, column=col), h)
        row += 1

        total = sum(b.get('count', 0) for b in by_building.values())

        for building, info in by_building.items():
            count = info.get('count', 0)
            ratio = round(count / total * 100, 1) if total > 0 else 0

            fault_types = info.get('fault_types', {})
            fault_str = '、'.join([f'{k}({v})' for k, v in sorted(fault_types.items(), key=lambda x: x[1], reverse=True)])

            status_counts = info.get('status_counts', {})
            resolved = status_counts.get('已处理', 0) + status_counts.get('已解决', 0)
            pending = status_counts.get('未处理', 0) + status_counts.get('处理中', 0)

            self._set_cell(ws.cell(row=row, column=1), building, bold=True)
            self._set_cell(ws.cell(row=row, column=2), count, align='center')
            self._set_cell(ws.cell(row=row, column=3), f'{ratio}%', align='center')
            self._set_cell(ws.cell(row=row, column=4), fault_str)
            self._set_cell(ws.cell(row=row, column=5), resolved, align='center')
            self._set_cell(ws.cell(row=row, column=6), pending, align='center')
            row += 1

        col_widths = [15, 10, 10, 40, 10, 10]
        for i, w in enumerate(col_widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = w

    def _excel_by_fault_type(self, wb, analysis: dict):
        """Sheet 5: 按故障类型统计"""
        from openpyxl.utils import get_column_letter

        ws = wb.create_sheet('按故障类型统计')

        by_fault = analysis.get('repairs_by_fault_type', {})

        ws.merge_cells('A1:E1')
        self._set_title(ws['A1'], '按故障类型统计')

        row = 3

        headers = ['故障类型', '报修次数', '占比', '主要楼栋', '主要原因']
        for col, h in enumerate(headers, 1):
            self._set_header(ws.cell(row=row, column=col), h)
        row += 1

        total = sum(f.get('count', 0) for f in by_fault.values())

        for ft, info in by_fault.items():
            count = info.get('count', 0)
            ratio = round(count / total * 100, 1) if total > 0 else 0

            buildings = info.get('buildings', {})
            top_buildings = sorted(buildings.items(), key=lambda x: x[1], reverse=True)[:3]
            building_str = '、'.join([f'{k}({v})' for k, v in top_buildings])

            causes = info.get('top_causes', {})
            top_causes = list(causes.keys())[:3]
            cause_str = '、'.join(top_causes)

            self._set_cell(ws.cell(row=row, column=1), ft, bold=True)
            self._set_cell(ws.cell(row=row, column=2), count, align='center')
            self._set_cell(ws.cell(row=row, column=3), f'{ratio}%', align='center')
            self._set_cell(ws.cell(row=row, column=4), building_str)
            self._set_cell(ws.cell(row=row, column=5), cause_str)
            row += 1

        col_widths = [12, 10, 10, 30, 40]
        for i, w in enumerate(col_widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = w

    def _excel_by_handler(self, wb, analysis: dict):
        """Sheet 6: 按处理人统计"""
        from openpyxl.utils import get_column_letter

        ws = wb.create_sheet('按处理人统计')

        by_handler = analysis.get('repairs_by_handler', {})

        ws.merge_cells('A1:D1')
        self._set_title(ws['A1'], '按处理人统计')

        row = 3

        headers = ['处理人', '处理次数', '已处理', '处理率']
        for col, h in enumerate(headers, 1):
            self._set_header(ws.cell(row=row, column=col), h)
        row += 1

        for handler, info in by_handler.items():
            self._set_cell(ws.cell(row=row, column=1), handler, bold=True)
            self._set_cell(ws.cell(row=row, column=2), info.get('count', 0), align='center')
            self._set_cell(ws.cell(row=row, column=3), info.get('resolved', 0), align='center')
            self._set_cell(ws.cell(row=row, column=4), f"{info.get('resolved_rate', 0)}%", align='center')
            row += 1

        col_widths = [15, 12, 12, 12]
        for i, w in enumerate(col_widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = w

    def _excel_by_college(self, wb, analysis: dict):
        """Sheet 7: 按学院统计"""
        from openpyxl.utils import get_column_letter

        ws = wb.create_sheet('按学院统计')

        by_college = analysis.get('repairs_by_college', {})

        ws.merge_cells('A1:D1')
        self._set_title(ws['A1'], '按学院统计')

        row = 3

        headers = ['学院', '报修次数', '外聘教师数', '占比']
        for col, h in enumerate(headers, 1):
            self._set_header(ws.cell(row=row, column=col), h)
        row += 1

        total = sum(c.get('count', 0) for c in by_college.values())

        for college, info in by_college.items():
            count = info.get('count', 0)
            ratio = round(count / total * 100, 1) if total > 0 else 0

            self._set_cell(ws.cell(row=row, column=1), college, bold=True)
            self._set_cell(ws.cell(row=row, column=2), count, align='center')
            self._set_cell(ws.cell(row=row, column=3), info.get('external_count', 0), align='center')
            self._set_cell(ws.cell(row=row, column=4), f'{ratio}%', align='center')
            row += 1

        col_widths = [25, 12, 12, 12]
        for i, w in enumerate(col_widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = w

    def _excel_week_comparison(self, wb, analysis: dict):
        """Sheet 8: 周对比分析"""
        from openpyxl.utils import get_column_letter

        ws = wb.create_sheet('周对比分析')

        comparison = analysis.get('week_comparison', {})

        ws.merge_cells('A1:E1')
        self._set_title(ws['A1'], '周对比分析')

        row = 3

        # 整体对比
        self._set_subtitle(ws.cell(row=row, column=1), '1. 整体指标对比')
        row += 1

        headers = ['指标', '本周', '上周', '变化量', '趋势']
        for col, h in enumerate(headers, 1):
            self._set_header(ws.cell(row=row, column=col), h)
        row += 1

        for item in comparison.get('overview_table', []):
            self._set_cell(ws.cell(row=row, column=1), item.get('metric', ''), bold=True)
            self._set_cell(ws.cell(row=row, column=2), item.get('this_week', ''), align='center')
            self._set_cell(ws.cell(row=row, column=3), item.get('last_week', ''), align='center')
            self._set_cell(ws.cell(row=row, column=4), item.get('change', ''), align='center')
            self._set_cell(ws.cell(row=row, column=5), item.get('trend', ''), align='center')
            row += 1

        row += 2

        # 故障类型对比
        self._set_subtitle(ws.cell(row=row, column=1), '2. 故障类型对比')
        row += 1

        headers = ['故障类型', '本周', '上周', '变化', '趋势', '需关注']
        for col, h in enumerate(headers, 1):
            self._set_header(ws.cell(row=row, column=col), h)
        row += 1

        for item in comparison.get('fault_type_comparison', []):
            self._set_cell(ws.cell(row=row, column=1), item.get('type', ''), bold=True)
            self._set_cell(ws.cell(row=row, column=2), item.get('this_week', ''), align='center')
            self._set_cell(ws.cell(row=row, column=3), item.get('last_week', ''), align='center')
            self._set_cell(ws.cell(row=row, column=4), item.get('change', ''), align='center')
            self._set_cell(ws.cell(row=row, column=5), item.get('trend', ''), align='center')
            self._set_cell(ws.cell(row=row, column=6), '⚠️ 是' if item.get('need_attention') else '否', align='center')
            row += 1

        row += 2

        # 楼栋对比
        self._set_subtitle(ws.cell(row=row, column=1), '3. 楼栋对比')
        row += 1

        headers = ['楼栋', '本周', '上周', '变化', '趋势', '需关注']
        for col, h in enumerate(headers, 1):
            self._set_header(ws.cell(row=row, column=col), h)
        row += 1

        for item in comparison.get('building_comparison', []):
            self._set_cell(ws.cell(row=row, column=1), item.get('building', ''), bold=True)
            self._set_cell(ws.cell(row=row, column=2), item.get('this_week', ''), align='center')
            self._set_cell(ws.cell(row=row, column=3), item.get('last_week', ''), align='center')
            self._set_cell(ws.cell(row=row, column=4), item.get('change', ''), align='center')
            self._set_cell(ws.cell(row=row, column=5), item.get('trend', ''), align='center')
            self._set_cell(ws.cell(row=row, column=6), '⚠️ 是' if item.get('need_attention') else '否', align='center')
            row += 1

        col_widths = [15, 12, 12, 12, 10, 10]
        for i, w in enumerate(col_widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = w

    def _excel_advice(self, wb, advice: list):
        """Sheet 9: 智能建议"""
        from openpyxl.utils import get_column_letter

        ws = wb.create_sheet('智能建议')

        ws.merge_cells('A1:D1')
        self._set_title(ws['A1'], '智能建议与改进措施')

        row = 3

        headers = ['级别', '分类', '问题描述', '建议措施']
        for col, h in enumerate(headers, 1):
            self._set_header(ws.cell(row=row, column=col), h)
        row += 1

        level_names = {'alert': '🚨 紧急', 'warning': '⚠️ 警告', 'info': 'ℹ️ 建议'}

        for item in advice:
            self._set_cell(ws.cell(row=row, column=1), level_names.get(item.get('level', ''), ''))
            self._set_cell(ws.cell(row=row, column=2), item.get('category', ''))
            self._set_cell(ws.cell(row=row, column=3), item.get('content', ''))
            self._set_cell(ws.cell(row=row, column=4), item.get('action', ''))
            row += 1

        col_widths = [12, 12, 50, 40]
        for i, w in enumerate(col_widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = w

    def _excel_multi_week_trend(self, wb, analysis: dict):
        """Sheet 10: 多周趋势"""
        from openpyxl.utils import get_column_letter

        ws = wb.create_sheet('多周趋势')
        trend = analysis.get('multi_week_trend', {})
        weeks = trend.get('weeks', [])

        ws.merge_cells('A1:G1')
        self._set_title(ws['A1'], '近4周趋势分析')

        row = 3
        self._set_subtitle(ws.cell(row=row, column=1), f'趋势总结：{trend.get("summary", "")}')
        row += 2

        # 表头
        headers = ['周次', '日期范围', '报修总量', '已处理', '待处理', '外聘教师', '设备更换']
        for col, h in enumerate(headers, 1):
            self._set_header(ws.cell(row=row, column=col), h)
        row += 1

        for w in weeks:
            self._set_cell(ws.cell(row=row, column=1), w.get('week', ''), bold=True)
            self._set_cell(ws.cell(row=row, column=2), w.get('date_range', ''), align='center')
            self._set_cell(ws.cell(row=row, column=3), w.get('total', 0), align='center')
            self._set_cell(ws.cell(row=row, column=4), w.get('resolved', 0), align='center')
            self._set_cell(ws.cell(row=row, column=5), w.get('pending', 0), align='center')
            self._set_cell(ws.cell(row=row, column=6), w.get('external_teacher', 0), align='center')
            self._set_cell(ws.cell(row=row, column=7), w.get('device_replace', 0), align='center')
            row += 1

        # 趋势信息
        row += 1
        self._set_subtitle(ws.cell(row=row, column=1), '趋势分析')
        row += 1
        self._set_cell(ws.cell(row=row, column=1), f'整体趋势：{trend.get("trend_label", "")}')
        row += 1
        self._set_cell(ws.cell(row=row, column=1), f'变化加速度：{trend.get("acceleration_label", "")}')

        col_widths = [12, 18, 12, 12, 12, 12, 12]
        for i, w in enumerate(col_widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = w

    def _excel_equipment_health(self, wb, analysis: dict):
        """Sheet 11: 设备健康度"""
        from openpyxl.utils import get_column_letter
        from openpyxl.styles import PatternFill

        ws = wb.create_sheet('设备健康度')
        health = analysis.get('equipment_health', {})
        scores = health.get('scores', [])

        ws.merge_cells('A1:J1')
        self._set_title(ws['A1'], '设备健康度评分')

        row = 2
        summary = health.get('summary', {})
        self._set_subtitle(ws.cell(row=row, column=1),
            f'全校平均分：{summary.get("avg_score", 0)} | 不健康教室：{summary.get("unhealthy_count", 0)}/{summary.get("total_classrooms", 0)}')
        row += 2

        # 表头
        headers = ['教室', '健康评分', '健康等级', '学期报修', '本周报修', '设备更换', '超时未处理', '主要故障类型', '重复故障', '扣分原因']
        for col, h in enumerate(headers, 1):
            self._set_header(ws.cell(row=row, column=col), h)
        row += 1

        # 等级颜色
        level_fills = {
            'danger': PatternFill(start_color='FFF0F0', end_color='FFF0F0', fill_type='solid'),
            'warning': PatternFill(start_color='FFF8E1', end_color='FFF8E1', fill_type='solid'),
            'attention': PatternFill(start_color='F0F7FF', end_color='F0F7FF', fill_type='solid'),
            'healthy': PatternFill(start_color='F0FFF0', end_color='F0FFF0', fill_type='solid'),
        }

        for item in scores[:30]:  # 最多显示30间
            score_val = item.get('health_score', 0)
            level = item.get('health_level', 'healthy')

            for col in range(1, 11):
                cell = ws.cell(row=row, column=col)
                if level in level_fills:
                    cell.fill = level_fills[level]

            self._set_cell(ws.cell(row=row, column=1), item.get('classroom', ''), bold=True)
            self._set_cell(ws.cell(row=row, column=2), score_val, bold=True, align='center')
            self._set_cell(ws.cell(row=row, column=3), item.get('health_label', ''), align='center')
            self._set_cell(ws.cell(row=row, column=4), item.get('semester_count', 0), align='center')
            self._set_cell(ws.cell(row=row, column=5), item.get('this_week_count', 0), align='center')
            self._set_cell(ws.cell(row=row, column=6), item.get('replace_count', 0), align='center')
            self._set_cell(ws.cell(row=row, column=7), item.get('overdue_count', 0), align='center')

            top_ft = item.get('top_fault_types', [])
            self._set_cell(ws.cell(row=row, column=8), '、'.join([f"{t['type']}({t['count']})" for t in top_ft[:2]]))

            repeat_ft = item.get('repeat_fault_types', [])
            self._set_cell(ws.cell(row=row, column=9), '、'.join([f"{t['type']}({t['count']})" for t in repeat_ft[:2]]))

            self._set_cell(ws.cell(row=row, column=10), '; '.join(item.get('deductions', [])[:3]))
            row += 1

        col_widths = [12, 10, 10, 10, 10, 10, 10, 25, 25, 40]
        for i, w in enumerate(col_widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = w

        ws.freeze_panes = 'A4'

    def _excel_building_profile(self, wb, analysis: dict):
        """Sheet 12: 楼栋画像"""
        from openpyxl.utils import get_column_letter

        ws = wb.create_sheet('楼栋画像')
        profile = analysis.get('building_profile', {})
        profiles = profile.get('profiles', {})
        school_avg = profile.get('school_average', {})

        ws.merge_cells('A1:J1')
        self._set_title(ws['A1'], '楼栋特征画像')

        row = 2
        self._set_subtitle(ws.cell(row=row, column=1),
            f'全校指标参考：外聘比{school_avg.get("external_ratio", 0)}% | 更换率{school_avg.get("replace_ratio", 0)}% | 平均处理{school_avg.get("avg_process_days", 0)}天')
        row += 2

        headers = ['楼栋', '学期报修', '本周报修', '典型故障模式', '外聘比', '更换率', '平均处理', '高峰时段', '教室数', '未处理']
        for col, h in enumerate(headers, 1):
            self._set_header(ws.cell(row=row, column=col), h)
        row += 1

        for building, info in profiles.items():
            # 标记外聘比/更换率/处理时间偏高的项
            ext_label = f"{info.get('external_ratio', 0)}%"
            if info.get('external_vs_avg') == 'high':
                ext_label += ' ⚠️'

            replace_label = f"{info.get('replace_ratio', 0)}%"
            if info.get('replace_vs_avg') == 'high':
                replace_label += ' ⚠️'

            process_label = f"{info.get('avg_process_days', 0)}天"
            if info.get('process_vs_avg') == 'high':
                process_label += ' ⚠️'

            self._set_cell(ws.cell(row=row, column=1), building, bold=True)
            self._set_cell(ws.cell(row=row, column=2), info.get('semester_total', 0), align='center')
            self._set_cell(ws.cell(row=row, column=3), info.get('this_week_count', 0), align='center')
            self._set_cell(ws.cell(row=row, column=4), info.get('typical_fault_label', ''))
            self._set_cell(ws.cell(row=row, column=5), ext_label, align='center')
            self._set_cell(ws.cell(row=row, column=6), replace_label, align='center')
            self._set_cell(ws.cell(row=row, column=7), process_label, align='center')
            self._set_cell(ws.cell(row=row, column=8), info.get('peak_section', ''), align='center')
            self._set_cell(ws.cell(row=row, column=9), info.get('classroom_count', 0), align='center')
            self._set_cell(ws.cell(row=row, column=10), info.get('unhandled_count', 0), align='center')
            row += 1

        col_widths = [12, 10, 10, 35, 12, 12, 12, 10, 8, 8]
        for i, w in enumerate(col_widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = w

    def _excel_response_time(self, wb, analysis: dict):
        """Sheet 13: 响应时效"""
        from openpyxl.utils import get_column_letter

        ws = wb.create_sheet('响应时效')
        rtm = analysis.get('response_time_matrix', {})

        ws.merge_cells('A1:H1')
        self._set_title(ws['A1'], '响应时效分析')

        # 按处理人
        row = 3
        self._set_subtitle(ws.cell(row=row, column=1), '一、按处理人时效')
        row += 1

        headers = ['处理人', '处理总量', '已处理', '完成率', '平均天数', '按时率(≤2天)', '涉及教室']
        for col, h in enumerate(headers, 1):
            self._set_header(ws.cell(row=row, column=col), h)
        row += 1

        for h in rtm.get('by_handler', []):
            self._set_cell(ws.cell(row=row, column=1), h.get('handler', ''), bold=True)
            self._set_cell(ws.cell(row=row, column=2), h.get('total', 0), align='center')
            self._set_cell(ws.cell(row=row, column=3), h.get('resolved', 0), align='center')
            self._set_cell(ws.cell(row=row, column=4), f"{h.get('completed_rate', 0)}%", align='center')
            self._set_cell(ws.cell(row=row, column=5), f"{h.get('avg_process_days', 0)}天", align='center')
            self._set_cell(ws.cell(row=row, column=6), f"{h.get('on_time_rate', 0)}%", align='center')
            self._set_cell(ws.cell(row=row, column=7), h.get('classroom_count', 0), align='center')
            row += 1

        # 按故障类型
        row += 2
        self._set_subtitle(ws.cell(row=row, column=1), '二、按故障类型时效')
        row += 1

        ft_headers = ['故障类型', '报修量', '已处理', '平均天数']
        for col, h in enumerate(ft_headers, 1):
            self._set_header(ws.cell(row=row, column=col), h)
        row += 1

        for ft in rtm.get('by_fault_type', []):
            self._set_cell(ws.cell(row=row, column=1), ft.get('fault_type', ''), bold=True)
            self._set_cell(ws.cell(row=row, column=2), ft.get('total', 0), align='center')
            self._set_cell(ws.cell(row=row, column=3), ft.get('resolved', 0), align='center')
            self._set_cell(ws.cell(row=row, column=4), f"{ft.get('avg_process_days', 0)}天", align='center')
            row += 1

        # 瓶颈
        bottlenecks = rtm.get('bottlenecks', [])
        if bottlenecks:
            row += 2
            self._set_subtitle(ws.cell(row=row, column=1), '三、响应瓶颈')
            row += 1
            for bn in bottlenecks:
                self._set_cell(ws.cell(row=row, column=1), f'⚠️ {bn.get("issue", "")}')
                row += 1
                self._set_cell(ws.cell(row=row, column=1), f'   💡 {bn.get("suggestion", "")}')
                row += 1

        col_widths = [15, 12, 12, 12, 14, 14, 10]
        for i, w in enumerate(col_widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = w

    def _excel_workload(self, wb, analysis: dict):
        """Sheet 14: 工作负载"""
        from openpyxl.utils import get_column_letter

        ws = wb.create_sheet('工作负载')
        workload = analysis.get('workload_analysis', {})

        ws.merge_cells('A1:G1')
        self._set_title(ws['A1'], '处理人工作负载分析')

        row = 3
        self._set_subtitle(ws.cell(row=row, column=1),
            f'负载评估：{workload.get("balance_assessment", "")}')
        row += 2

        headers = ['处理人', '学期总量', '本周量', '已处理', '完成率', '待处理积压', '日均处理']
        for col, h in enumerate(headers, 1):
            self._set_header(ws.cell(row=row, column=col), h)
        row += 1

        for h in workload.get('handlers', []):
            self._set_cell(ws.cell(row=row, column=1), h.get('handler', ''), bold=True)
            self._set_cell(ws.cell(row=row, column=2), h.get('semester_total', 0), align='center')
            self._set_cell(ws.cell(row=row, column=3), h.get('week_total', 0), align='center')
            self._set_cell(ws.cell(row=row, column=4), h.get('resolved', 0), align='center')
            self._set_cell(ws.cell(row=row, column=5), f"{h.get('completed_rate', 0)}%", align='center')
            self._set_cell(ws.cell(row=row, column=6), h.get('pending', 0), align='center')
            self._set_cell(ws.cell(row=row, column=7), f"{h.get('daily_avg', 0):.2f}件/天", align='center')
            row += 1

        col_widths = [15, 12, 10, 10, 10, 12, 14]
        for i, w in enumerate(col_widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = w

    # ============================================================
    # Word 报告
    # ============================================================

    def render_word(self, analysis: dict, advice: list = None) -> BytesIO:
        """渲染完整 Word 报告"""
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = Document()

        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style.font.size = Pt(10.5)
        style.paragraph_format.line_spacing = 1.5

        report_type = analysis.get('report_type', '')
        date_range = analysis.get('date_range', {})
        overview = analysis.get('overview', {})

        # ============ 封面 ============
        for _ in range(3):
            doc.add_paragraph()

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('重庆移通学院綦江校区')
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(47, 84, 150)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f'多媒体设备运维{self._get_type_name(report_type)}')
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(47, 84, 150)

        doc.add_paragraph()

        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = info.add_run(date_range.get('label', ''))
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_paragraph()

        time_p = doc.add_paragraph()
        time_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = time_p.add_run(f'报告生成时间：{datetime.now().strftime("%Y年%m月%d日")}')
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(150, 150, 150)

        doc.add_page_break()

        # ============ 一、核心指标 ============
        doc.add_heading('一、核心指标概览', level=1)

        # 核心指标表格
        table = doc.add_table(rows=7, cols=2)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        vs = overview.get('vs_last_week', {})
        metrics = [
            ('报修总量', f"{overview.get('total_count', 0)}件（较上周{vs.get('total_change', 0):+d}件，{vs.get('total_change_rate', 0):+.1f}%）"),
            ('已处理', f"{overview.get('resolved_count', 0)}件"),
            ('待处理', f"{overview.get('pending_count', 0)}件（处理中{overview.get('processing_count', 0)}件，未处理{overview.get('unhandled_count', 0)}件）"),
            ('处理率', f"{overview.get('resolved_rate', 0)}%（较上周{vs.get('rate_change', 0):+.1f}%）"),
            ('外聘教师报修', f"{overview.get('external_teacher_count', 0)}件"),
            ('设备更换', f"{overview.get('device_replace_count', 0)}次"),
        ]

        table.rows[0].cells[0].text = '指标'
        table.rows[0].cells[1].text = '数值'

        for i, (label, value) in enumerate(metrics):
            table.rows[i + 1].cells[0].text = label
            table.rows[i + 1].cells[1].text = value

        doc.add_paragraph()

        # ============ 二、报修明细 ============
        doc.add_heading('二、报修记录明细', level=1)

        records = analysis.get('records_detail', [])
        if records:
            doc.add_paragraph(f'本周共有 {len(records)} 条报修记录，详见下表：')

            table = doc.add_table(rows=len(records) + 1, cols=8)
            table.style = 'Light Grid Accent 1'

            headers = ['日期', '教室', '故障类型', '故障原因', '状态', '报修人', '处理人', '处理方案']
            for j, h in enumerate(headers):
                table.rows[0].cells[j].text = h

            for i, r in enumerate(records):
                row = table.rows[i + 1]
                row.cells[0].text = f"{r.get('date', '')} {r.get('weekday', '')}"
                row.cells[1].text = r.get('classroom', '')
                row.cells[2].text = r.get('fault_type', '')
                row.cells[3].text = r.get('fault_cause', '')[:20] + ('...' if len(r.get('fault_cause', '')) > 20 else '')
                row.cells[4].text = r.get('status', '')
                row.cells[5].text = r.get('reporter_name', '')
                row.cells[6].text = r.get('handler_name', '')
                row.cells[7].text = r.get('solution', '')[:20] + ('...' if len(r.get('solution', '')) > 20 else '')
        else:
            doc.add_paragraph('本周暂无报修记录。')

        doc.add_paragraph()

        # ============ 三、按教室统计 ============
        doc.add_heading('三、按教室统计', level=1)

        by_classroom = analysis.get('repairs_by_classroom', {})
        if by_classroom:
            table = doc.add_table(rows=min(len(by_classroom), 10) + 1, cols=4)
            table.style = 'Light Grid Accent 1'

            table.rows[0].cells[0].text = '教室'
            table.rows[0].cells[1].text = '本周次数'
            table.rows[0].cells[2].text = '学期累计'
            table.rows[0].cells[3].text = '故障类型'

            for i, (classroom, info) in enumerate(list(by_classroom.items())[:10]):
                row = table.rows[i + 1]
                repairs = info.get('repairs', [])
                fault_types = set(r.get('fault_type', '') for r in repairs if r.get('fault_type'))

                row.cells[0].text = classroom
                row.cells[1].text = str(info.get('count', 0))
                row.cells[2].text = str(info.get('semester_total', 0))
                row.cells[3].text = '、'.join(fault_types)

                if info.get('is_frequent'):
                    row.cells[0].text += ' ⚠️'

        doc.add_paragraph()

        # ============ 四、按楼栋统计 ============
        doc.add_heading('四、按楼栋统计', level=1)

        by_building = analysis.get('repairs_by_building', {})
        if by_building:
            table = doc.add_table(rows=len(by_building) + 1, cols=3)
            table.style = 'Light Grid Accent 1'

            table.rows[0].cells[0].text = '楼栋'
            table.rows[0].cells[1].text = '报修次数'
            table.rows[0].cells[2].text = '占比'

            total = sum(b.get('count', 0) for b in by_building.values())
            for i, (building, info) in enumerate(by_building.items()):
                row = table.rows[i + 1]
                count = info.get('count', 0)
                ratio = round(count / total * 100, 1) if total > 0 else 0
                row.cells[0].text = building
                row.cells[1].text = str(count)
                row.cells[2].text = f'{ratio}%'

        doc.add_paragraph()

        # ============ 五、按故障类型统计 ============
        doc.add_heading('五、按故障类型统计', level=1)

        by_fault = analysis.get('repairs_by_fault_type', {})
        if by_fault:
            table = doc.add_table(rows=len(by_fault) + 1, cols=3)
            table.style = 'Light Grid Accent 1'

            table.rows[0].cells[0].text = '故障类型'
            table.rows[0].cells[1].text = '报修次数'
            table.rows[0].cells[2].text = '占比'

            total = sum(f.get('count', 0) for f in by_fault.values())
            for i, (ft, info) in enumerate(by_fault.items()):
                row = table.rows[i + 1]
                count = info.get('count', 0)
                ratio = round(count / total * 100, 1) if total > 0 else 0
                row.cells[0].text = ft
                row.cells[1].text = str(count)
                row.cells[2].text = f'{ratio}%'

        doc.add_paragraph()

        # ============ 六、按学院统计 ============
        doc.add_heading('六、按学院统计', level=1)

        by_college = analysis.get('repairs_by_college', {})
        if by_college:
            table = doc.add_table(rows=min(len(by_college), 10) + 1, cols=3)
            table.style = 'Light Grid Accent 1'

            table.rows[0].cells[0].text = '学院'
            table.rows[0].cells[1].text = '报修次数'
            table.rows[0].cells[2].text = '外聘教师数'

            for i, (college, info) in enumerate(list(by_college.items())[:10]):
                row = table.rows[i + 1]
                row.cells[0].text = college
                row.cells[1].text = str(info.get('count', 0))
                row.cells[2].text = str(info.get('external_count', 0))

        doc.add_paragraph()

        # ============ 七、按处理人统计 ============
        doc.add_heading('七、按处理人统计', level=1)

        by_handler = analysis.get('repairs_by_handler', {})
        if by_handler:
            table = doc.add_table(rows=len(by_handler) + 1, cols=3)
            table.style = 'Light Grid Accent 1'

            table.rows[0].cells[0].text = '处理人'
            table.rows[0].cells[1].text = '处理次数'
            table.rows[0].cells[2].text = '处理率'

            for i, (handler, info) in enumerate(by_handler.items()):
                row = table.rows[i + 1]
                row.cells[0].text = handler
                row.cells[1].text = str(info.get('count', 0))
                row.cells[2].text = f"{info.get('resolved_rate', 0)}%"

        doc.add_paragraph()

        # ============ 多周趋势分析（仅周报） ============
        if report_type == 'weekly':
            self._word_multi_week_trend(doc, analysis)

        # ============ 设备健康度 ============
        self._word_equipment_health(doc, analysis)

        # ============ 楼栋画像 ============
        self._word_building_profile(doc, analysis)

        # ============ 响应时效 ============
        self._word_response_time(doc, analysis)

        # ============ 工作负载 ============
        self._word_workload(doc, analysis)

        # ============ 八、智能建议 ============
        doc.add_heading('八、智能建议与改进措施', level=1)

        if advice:
            for i, item in enumerate(advice, 1):
                level = item.get('level', 'info')
                level_text = {'alert': '紧急', 'warning': '警告', 'info': '建议'}.get(level, '')

                p = doc.add_paragraph()
                run = p.add_run(f'{i}. [{level_text}] {item.get("title", "")}')
                run.bold = True
                run.font.size = Pt(11)

                if level == 'alert':
                    run.font.color.rgb = RGBColor(239, 68, 68)
                elif level == 'warning':
                    run.font.color.rgb = RGBColor(245, 158, 11)

                p = doc.add_paragraph(f'问题：{item.get("content", "")}')
                p.paragraph_format.left_indent = Cm(1)

                p = doc.add_paragraph(f'建议：{item.get("action", "")}')
                p.paragraph_format.left_indent = Cm(1)

                doc.add_paragraph()
        else:
            doc.add_paragraph('暂无特别建议。')

        # 保存
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _word_multi_week_trend(self, doc, analysis: dict):
        """Word: 多周趋势分析"""
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        trend = analysis.get('multi_week_trend', {})
        weeks = trend.get('weeks', [])

        doc.add_heading('多周趋势分析', level=1)

        doc.add_paragraph(f'趋势判断：{trend.get("summary", "")}')
        doc.add_paragraph()

        if weeks:
            table = doc.add_table(rows=len(weeks) + 1, cols=7)
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            headers = ['周次', '日期', '总量', '已处理', '待处理', '外聘教师', '设备更换']
            for j, h in enumerate(headers):
                table.rows[0].cells[j].text = h

            for i, w in enumerate(weeks):
                row = table.rows[i + 1]
                row.cells[0].text = w.get('week', '')
                row.cells[1].text = w.get('date_range', '')
                row.cells[2].text = str(w.get('total', 0))
                row.cells[3].text = str(w.get('resolved', 0))
                row.cells[4].text = str(w.get('pending', 0))
                row.cells[5].text = str(w.get('external_teacher', 0))
                row.cells[6].text = str(w.get('device_replace', 0))

        doc.add_paragraph()

    def _word_equipment_health(self, doc, analysis: dict):
        """Word: 设备健康度"""
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.table import WD_TABLE_ALIGNMENT

        health = analysis.get('equipment_health', {})
        scores = health.get('scores', [])
        summary = health.get('summary', {})
        problems = health.get('this_week_problems', [])

        doc.add_heading('设备健康度报告', level=1)

        doc.add_paragraph(f'全校设备健康概览：平均分 {summary.get("avg_score", 0)} 分，共 {summary.get("total_classrooms", 0)} 间教室纳入评估，其中 {summary.get("unhealthy_count", 0)} 间处于"警告"或"危险"级别。')
        doc.add_paragraph()

        if problems:
            doc.add_paragraph('本周重点关注教室：')
            for item in problems[:5]:
                classroom = item.get('classroom', '')
                score = item.get('health_score', 0)
                label = item.get('health_label', '')
                semester_count = item.get('semester_count', 0)
                repeat_faults = item.get('repeat_fault_types', [])
                deductions = item.get('deductions', [])

                p = doc.add_paragraph()
                run = p.add_run(f'• {classroom}（{score}分/{label}）')
                run.bold = True
                if score < 40:
                    run.font.color.rgb = RGBColor(239, 68, 68)
                elif score < 60:
                    run.font.color.rgb = RGBColor(245, 158, 11)

                p = doc.add_paragraph(f'  学期报修 {semester_count} 次，扣分项：{"；".join(deductions[:3])}')
                p.paragraph_format.left_indent = Cm(1)
        else:
            doc.add_paragraph('本周暂无设备健康度预警。')

        doc.add_paragraph()

        # 健康度排行表（Top 10 最差）
        worst = scores[:10]
        if worst:
            table = doc.add_table(rows=min(len(worst), 10) + 1, cols=5)
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            table.rows[0].cells[0].text = '教室'
            table.rows[0].cells[1].text = '评分'
            table.rows[0].cells[2].text = '等级'
            table.rows[0].cells[3].text = '学期次数'
            table.rows[0].cells[4].text = '主要故障'

            for i, item in enumerate(worst[:10]):
                row = table.rows[i + 1]
                row.cells[0].text = item.get('classroom', '')
                row.cells[1].text = str(item.get('health_score', 0))
                row.cells[2].text = item.get('health_label', '')
                row.cells[3].text = str(item.get('semester_count', 0))
                top_ft = item.get('top_fault_types', [])
                row.cells[4].text = '、'.join([f"{t['type']}({t['count']})" for t in top_ft[:2]])

        doc.add_paragraph()

    def _word_building_profile(self, doc, analysis: dict):
        """Word: 楼栋画像"""
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.table import WD_TABLE_ALIGNMENT

        profile = analysis.get('building_profile', {})
        profiles = profile.get('profiles', {})
        attention = profile.get('attention_buildings', [])
        school_avg = profile.get('school_average', {})

        doc.add_heading('楼栋特征分析', level=1)

        doc.add_paragraph(f'全校参考值：外聘比 {school_avg.get("external_ratio", 0)}%，更换率 {school_avg.get("replace_ratio", 0)}%，平均处理 {school_avg.get("avg_process_days", 0)} 天。')
        doc.add_paragraph()

        if profiles:
            table = doc.add_table(rows=len(profiles) + 1, cols=6)
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            table.rows[0].cells[0].text = '楼栋'
            table.rows[0].cells[1].text = '典型故障模式'
            table.rows[0].cells[2].text = '外聘比'
            table.rows[0].cells[3].text = '更换率'
            table.rows[0].cells[4].text = '平均处理'
            table.rows[0].cells[5].text = '高峰时段'

            for i, (building, info) in enumerate(profiles.items()):
                row = table.rows[i + 1]
                row.cells[0].text = building
                row.cells[1].text = info.get('typical_fault_label', '')
                row.cells[2].text = f"{info.get('external_ratio', 0)}%{' ⚠️' if info.get('external_vs_avg') == 'high' else ''}"
                row.cells[3].text = f"{info.get('replace_ratio', 0)}%{' ⚠️' if info.get('replace_vs_avg') == 'high' else ''}"
                row.cells[4].text = f"{info.get('avg_process_days', 0)}天{' ⚠️' if info.get('process_vs_avg') == 'high' else ''}"
                row.cells[5].text = info.get('peak_section', '')

        if attention:
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run(f'⚠️ 需要重点关注的楼栋：{"、".join(attention)}')
            run.bold = True
            run.font.color.rgb = RGBColor(245, 158, 11)

        doc.add_paragraph()

    def _word_response_time(self, doc, analysis: dict):
        """Word: 响应时效"""
        from docx.shared import Pt, RGBColor
        from docx.enum.table import WD_TABLE_ALIGNMENT

        rtm = analysis.get('response_time_matrix', {})

        doc.add_heading('响应时效分析', level=1)

        # 按处理人
        by_handler = rtm.get('by_handler', [])
        if by_handler:
            doc.add_paragraph('按处理人时效：')

            table = doc.add_table(rows=len(by_handler) + 1, cols=5)
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            table.rows[0].cells[0].text = '处理人'
            table.rows[0].cells[1].text = '总量'
            table.rows[0].cells[2].text = '完成率'
            table.rows[0].cells[3].text = '平均天数'
            table.rows[0].cells[4].text = '按时率(≤2天)'

            for i, h in enumerate(by_handler):
                row = table.rows[i + 1]
                row.cells[0].text = h.get('handler', '')
                row.cells[1].text = str(h.get('total', 0))
                row.cells[2].text = f"{h.get('completed_rate', 0)}%"
                row.cells[3].text = f"{h.get('avg_process_days', 0)}天"
                row.cells[4].text = f"{h.get('on_time_rate', 0)}%"

        doc.add_paragraph()

        # 瓶颈
        bottlenecks = rtm.get('bottlenecks', [])
        if bottlenecks:
            doc.add_paragraph('响应瓶颈：')
            for bn in bottlenecks:
                p = doc.add_paragraph()
                run = p.add_run(f'• {bn.get("issue", "")}')
                run.bold = True
                p = doc.add_paragraph(f'  建议：{bn.get("suggestion", "")}')

        doc.add_paragraph()

    def _word_workload(self, doc, analysis: dict):
        """Word: 工作负载"""
        from docx.shared import Pt, RGBColor
        from docx.enum.table import WD_TABLE_ALIGNMENT

        workload = analysis.get('workload_analysis', {})

        doc.add_heading('工作负载分析', level=1)

        doc.add_paragraph(f'负载评估：{workload.get("balance_assessment", "")}')
        doc.add_paragraph()

        handlers = workload.get('handlers', [])
        if handlers:
            table = doc.add_table(rows=len(handlers) + 1, cols=5)
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            table.rows[0].cells[0].text = '处理人'
            table.rows[0].cells[1].text = '学期总量'
            table.rows[0].cells[2].text = '完成率'
            table.rows[0].cells[3].text = '待处理积压'
            table.rows[0].cells[4].text = '日均处理'

            for i, h in enumerate(handlers):
                row = table.rows[i + 1]
                row.cells[0].text = h.get('handler', '')
                row.cells[1].text = str(h.get('semester_total', 0))
                row.cells[2].text = f"{h.get('completed_rate', 0)}%"
                row.cells[3].text = str(h.get('pending', 0))
                row.cells[4].text = f"{h.get('daily_avg', 0):.2f}件/天"

        if not workload.get('is_balanced', True):
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run('⚠️ 工作负载不均，建议合理调配任务分配。')
            run.bold = True
            run.font.color.rgb = RGBColor(245, 158, 11)

        doc.add_paragraph()

    def _get_type_name(self, report_type: str) -> str:
        return {'weekly': '周报', 'monthly': '月报', 'semester': '学期报告'}.get(report_type, '报告')


# 创建全局实例
renderer = ReportRenderer()
